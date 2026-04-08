"""
GuideChuna 평가모드 데이터 지표 정의서 - Word/PDF 생성 스크립트
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from fpdf import FPDF

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# 공통 데이터
# ============================================================

SESSION_INFO = [
    ("세션 ID", "평가 고유 식별자 (UUID)", "a1b2c3d4-..."),
    ("사용자", "로그인 사용자명 / ID", "홍길동 / 1023"),
    ("시나리오", "수행 시나리오명", "상부승모근, 대흉근"),
    ("모드", "학습 / 평가", "평가"),
    ("난이도", "초급 / 중급 / 상급 / 평가", "평가"),
    ("공식 평가 여부", "공식 점수 기록 대상인지", "O / X"),
    ("모의평가 여부", "상급자 모드의 모의평가인지", "O / X"),
    ("시도 횟수", "같은 시나리오 몇 회차 수행", "3"),
    ("시작/종료 시간", "평가 시작~종료 시각", "2026-04-03 15:30:00"),
    ("총 수행 시간", "전체 소요 시간 (초)", "245.3초 (4:05)"),
]

OVERALL_SCORE = [
    ("종합 점수", "전체 Step finalScore의 평균", "0 ~ 100"),
    ("종합 등급", "점수 기반 등급", "S / A+ / A / B+ / B / C+ / C / D / F"),
    ("전체 평균 유사도", "Phase별 평균 유사도의 평균", "0.0 ~ 1.0 (0% ~ 100%)"),
]

GRADE_TABLE = [
    ("S", "95점 이상"),
    ("A+", "90 ~ 94"),
    ("A", "85 ~ 89"),
    ("B+", "80 ~ 84"),
    ("B", "75 ~ 79"),
    ("C+", "70 ~ 74"),
    ("C", "65 ~ 69"),
    ("D", "60 ~ 64"),
    ("F", "60점 미만"),
]

SCORE_FORMULA = [
    ("유사도 점수", "60점 만점", "평균유사도 x 60"),
    ("안전성 점수", "40점 만점 (감점제)", "40점에서 아래 감점 적용"),
    ("  - 리밋 초과 진입", "회당 -2점", "위험 범위 진입 횟수"),
    ("  - 경고 구간 체류", "초당 -0.5점", "Warning 상태 누적 시간"),
    ("  - 위험 구간 체류", "초당 -1점", "Danger 상태 누적 시간"),
    ("  - 초과 구간 체류", "초당 -3점", "Exceeded 상태 누적 시간"),
]

SIMILARITY_METRICS = [
    ("평균 유사도", "가중 평균 (주동수 70% + 보조수 30%)", "0.0 ~ 1.0"),
    ("왼손 평균 유사도", "왼손 단독 평균", "0.0 ~ 1.0"),
    ("오른손 평균 유사도", "오른손 단독 평균", "0.0 ~ 1.0"),
    ("유사도 표준편차", "낮을수록 안정적으로 유지", "0.0 ~"),
    ("최저 유사도", "수행 중 가장 낮았던 순간", "0.0 ~ 1.0"),
    ("최고 유사도", "수행 중 가장 높았던 순간", "0.0 ~ 1.0"),
]

SAFETY_METRICS = [
    ("리밋 초과 진입 횟수", "위험 범위에 진입한 횟수 (상태 전환 시만 카운트)", "회"),
    ("경고 구간 누적 시간", "Warning 상태에 머문 시간", "초"),
    ("초과 구간 누적 시간", "Exceeded 상태에 머문 시간", "초"),
    ("최대 초과 비율", "가장 크게 넘어간 순간의 비율", "0.0 ~"),
    ("경고 횟수", "제한 범위 초과 경고 누적", "회"),
]

COMPLETION_METRICS = [
    ("완료 상태", "Step의 전체 수행 결과", "O(완료) / △(일부) / X(스킵)"),
    ("전체 SubStep 수", "해당 Step의 세부 단계 수", "정수"),
    ("완료된 SubStep", "정상 수행 완료한 수", "정수"),
    ("스킵된 SubStep", "시간 초과로 건너뛴 수", "정수"),
    ("소요 시간", "해당 Step에 걸린 시간", "초"),
    ("최종 점수", "유사도(60%) + 안전성(40%)", "0 ~ 100"),
    ("등급", "점수 기반 등급", "S ~ F"),
]

TIMELINE_METRICS = [
    ("경과 시간", "평가 시작 후 경과 (초)"),
    ("왼손 유사도", "해당 시점의 왼손 유사도"),
    ("오른손 유사도", "해당 시점의 오른손 유사도"),
    ("왼손 리밋 상태", "Normal / Warning / Danger / Exceeded"),
    ("오른손 리밋 상태", "Normal / Warning / Danger / Exceeded"),
    ("왼/오른손 프레임 비율", "가동범위 대비 현재 위치 비율"),
    ("왼/오른손 위치", "3D 좌표 (Vector3)"),
]

HIERARCHY_TEXT = """TrainingResultData (1 세션)
+-- 기본 정보 (사용자, 시나리오, 모드, 시간 등)
+-- 종합 통계 (점수, 등급, 유사도, 경고/초과/스킵 합산)
|
+-- PhaseResult[] (전부 / 중부 / 후부)
    +-- Phase 평균 유사도
    +-- Phase 총 시간
    +-- Phase 경고 횟수
    |
    +-- StepResult[] (제한장벽 확인 / 스트레칭 / 재평가 등)
        +-- 유사도 지표 6개 (평균, 좌/우, 표준편차, 최저, 최고)
        +-- 안전성 지표 5개 (초과횟수, 경고시간, 초과시간, 최대비율, 경고수)
        +-- 수행 완료도 6개 (상태, SubStep 수/완료/스킵, 시간, 점수/등급)
        +-- 시계열 데이터 (시간, 좌/우 유사도)"""


# ============================================================
# WORD (.docx)
# ============================================================

def set_cell_shading(cell, color_hex):
    """셀 배경색 설정"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    """헤더 강조 테이블 추가"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2F5496")

    # 데이터
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "D9E2F3")

    return table


def generate_docx():
    doc = Document()

    # 기본 폰트
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # 제목
    title = doc.add_heading('GuideChuna 평가모드 데이터 지표 정의서', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('VR 추나요법 교육 시스템 - 평가 결과 데이터 구조 및 산출 기준', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. 세션 기본 정보
    doc.add_heading('1. 세션 기본 정보', level=1)
    add_styled_table(doc, ["항목", "설명", "예시"], SESSION_INFO)

    # 2. 종합 점수
    doc.add_heading('2. 종합 점수 (전체 Step 평균)', level=1)
    add_styled_table(doc, ["항목", "설명", "범위"], OVERALL_SCORE)

    doc.add_heading('등급 기준', level=2)
    add_styled_table(doc, ["등급", "점수 범위"], GRADE_TABLE)

    # 3. 점수 산출 공식
    doc.add_heading('3. 최종 점수 산출 공식 (Step 단위)', level=1)
    p = doc.add_paragraph()
    run = p.add_run('최종 점수 = 유사도 점수(60%) + 안전성 점수(40%)')
    run.bold = True
    run.font.size = Pt(11)

    add_styled_table(doc, ["구성 요소", "배점", "산출 방식"], SCORE_FORMULA)
    doc.add_paragraph('* 최종 점수는 0~100 범위로 클램핑', style='List Bullet')

    # 4. Step별 상세 지표
    doc.add_heading('4. Step별 상세 지표', level=1)

    doc.add_heading('4-1. 유사도 (손 자세 정확도)', level=2)
    add_styled_table(doc, ["지표", "설명", "단위"], SIMILARITY_METRICS)
    doc.add_paragraph('유사도 구성: 조인트 로컬 포즈(20%) + 손바닥 위치(40%) + 손바닥 회전(40%)', style='List Bullet')

    doc.add_heading('4-2. 안전성 (가동범위 준수)', level=2)
    add_styled_table(doc, ["지표", "설명", "단위"], SAFETY_METRICS)

    doc.add_heading('4-3. 수행 완료도', level=2)
    add_styled_table(doc, ["지표", "설명", "값"], COMPLETION_METRICS)

    # 5. 시계열 데이터
    doc.add_heading('5. 시계열 데이터 (그래프용)', level=1)
    doc.add_paragraph('평가 중 일정 간격으로 기록되는 스냅샷:')
    add_styled_table(doc, ["항목", "설명"], TIMELINE_METRICS)
    doc.add_paragraph('CSV 저장 시에는 경과시간 + 좌/우 유사도만 저장 (경량화)', style='List Bullet')

    # 6. 데이터 계층 구조
    doc.add_heading('6. 데이터 계층 구조', level=1)
    p = doc.add_paragraph()
    run = p.add_run(HIERARCHY_TEXT)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # 7. CSV 저장 형식
    doc.add_heading('7. CSV 저장 형식', level=1)
    doc.add_paragraph('시나리오 완료 시 자동으로 2개의 CSV 파일 생성:')

    doc.add_heading('파일명 규칙', level=2)
    p = doc.add_paragraph()
    run = p.add_run('{날짜}_{시간}_{사용자}_{시나리오}_{세션ID}_{종류}.csv')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    doc.add_paragraph('예: 20260403_153000_홍길동_상부승모근_a1b2c3d4_summary.csv')

    doc.add_heading('요약 CSV (summary)', level=2)
    doc.add_paragraph('세션 정보 + Phase/Step별 1행. 34개 컬럼.')

    doc.add_heading('시계열 CSV (timeline)', level=2)
    doc.add_paragraph('시간축 유사도 변화 (웹뷰 그래프용). 6개 컬럼: SessionId, Phase, Step, Time, LeftSimilarity, RightSimilarity')

    doc.add_heading('저장 위치', level=2)
    doc.add_paragraph('에디터: Assets/Results/', style='List Bullet')
    doc.add_paragraph('빌드 (Quest): Application.persistentDataPath/Results/', style='List Bullet')

    path = os.path.join(OUTPUT_DIR, "GuideChuna_평가지표_정의서.docx")
    doc.save(path)
    print(f"Word saved: {path}")
    return path


# ============================================================
# PDF
# ============================================================

class KoreanPDF(FPDF):
    def __init__(self):
        super().__init__()
        # 한글 폰트 등록
        font_path = r"C:\Windows\Fonts\malgun.ttf"
        bold_path = r"C:\Windows\Fonts\malgunbd.ttf"
        self.add_font("malgun", "", font_path, uni=True)
        self.add_font("malgun", "B", bold_path, uni=True)
        self.set_auto_page_break(auto=True, margin=15)

    def header(self):
        self.set_font("malgun", "B", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 5, "GuideChuna 평가모드 데이터 지표 정의서", align="R", new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def footer(self):
        self.set_y(-15)
        self.set_font("malgun", "", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"- {self.page_no()} -", align="C")

    def section_title(self, title):
        self.set_font("malgun", "B", 13)
        self.set_text_color(47, 84, 150)
        self.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(47, 84, 150)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(3)

    def sub_title(self, title):
        self.set_font("malgun", "B", 11)
        self.set_text_color(68, 114, 196)
        self.cell(0, 8, title, new_x="LMARGIN", new_y="NEXT")
        self.ln(1)

    def body_text(self, text):
        self.set_font("malgun", "", 9)
        self.set_text_color(0, 0, 0)
        self.multi_cell(0, 5, text)
        self.ln(1)

    def note_text(self, text):
        self.set_font("malgun", "", 8)
        self.set_text_color(100, 100, 100)
        self.cell(5)
        self.multi_cell(0, 4.5, f"* {text}")
        self.ln(1)

    def add_table(self, headers, rows, col_widths=None):
        if col_widths is None:
            available = self.w - self.l_margin - self.r_margin
            col_widths = [available / len(headers)] * len(headers)

        # 헤더
        self.set_font("malgun", "B", 8)
        self.set_fill_color(47, 84, 150)
        self.set_text_color(255, 255, 255)
        for i, h in enumerate(headers):
            self.cell(col_widths[i], 7, h, border=1, fill=True, align="C")
        self.ln()

        # 데이터
        self.set_font("malgun", "", 8)
        self.set_text_color(0, 0, 0)
        for ri, row in enumerate(rows):
            if ri % 2 == 1:
                self.set_fill_color(217, 226, 243)
                fill = True
            else:
                fill = False

            max_lines = 1
            for ci, val in enumerate(row):
                lines = self.multi_cell(col_widths[ci], 5.5, str(val), border=0, split_only=True)
                max_lines = max(max_lines, len(lines))

            row_h = max_lines * 5.5

            # 페이지 넘김 체크
            if self.get_y() + row_h > self.h - self.b_margin:
                self.add_page()
                # 헤더 다시 그리기
                self.set_font("malgun", "B", 8)
                self.set_fill_color(47, 84, 150)
                self.set_text_color(255, 255, 255)
                for i, h in enumerate(headers):
                    self.cell(col_widths[i], 7, h, border=1, fill=True, align="C")
                self.ln()
                self.set_font("malgun", "", 8)
                self.set_text_color(0, 0, 0)

            if fill:
                self.set_fill_color(217, 226, 243)

            x_start = self.get_x()
            y_start = self.get_y()

            for ci, val in enumerate(row):
                self.set_xy(x_start + sum(col_widths[:ci]), y_start)
                self.cell(col_widths[ci], row_h, str(val), border=1, fill=fill)

            self.set_xy(x_start, y_start + row_h)
        self.ln(2)


def generate_pdf():
    pdf = KoreanPDF()
    pdf.add_page()

    # 표지
    pdf.ln(30)
    pdf.set_font("malgun", "B", 24)
    pdf.set_text_color(47, 84, 150)
    pdf.cell(0, 15, "GuideChuna", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("malgun", "B", 18)
    pdf.cell(0, 12, "평가모드 데이터 지표 정의서", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)
    pdf.set_font("malgun", "", 11)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 8, "VR 추나요법 교육 시스템", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 8, "평가 결과 데이터 구조 및 산출 기준", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(30)
    pdf.set_font("malgun", "", 10)
    pdf.set_text_color(130, 130, 130)
    pdf.cell(0, 8, "2026-04-03", align="C", new_x="LMARGIN", new_y="NEXT")

    # 1. 세션 기본 정보
    pdf.add_page()
    pdf.section_title("1. 세션 기본 정보")
    w = pdf.w - pdf.l_margin - pdf.r_margin
    pdf.add_table(["항목", "설명", "예시"], SESSION_INFO, [w*0.2, w*0.45, w*0.35])

    # 2. 종합 점수
    pdf.section_title("2. 종합 점수 (전체 Step 평균)")
    pdf.add_table(["항목", "설명", "범위"], OVERALL_SCORE, [w*0.2, w*0.45, w*0.35])

    pdf.sub_title("등급 기준")
    pdf.add_table(["등급", "점수 범위"], GRADE_TABLE, [w*0.3, w*0.7])

    # 3. 점수 산출 공식
    pdf.add_page()
    pdf.section_title("3. 최종 점수 산출 공식 (Step 단위)")
    pdf.set_font("malgun", "B", 11)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 8, "최종 점수 = 유사도 점수(60%) + 안전성 점수(40%)", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)
    pdf.add_table(["구성 요소", "배점", "산출 방식"], SCORE_FORMULA, [w*0.3, w*0.25, w*0.45])
    pdf.note_text("최종 점수는 0~100 범위로 클램핑")

    # 4. Step별 상세 지표
    pdf.section_title("4. Step별 상세 지표")

    pdf.sub_title("4-1. 유사도 (손 자세 정확도)")
    pdf.add_table(["지표", "설명", "단위"], SIMILARITY_METRICS, [w*0.25, w*0.5, w*0.25])
    pdf.note_text("유사도 구성: 조인트 로컬 포즈(20%) + 손바닥 위치(40%) + 손바닥 회전(40%)")

    pdf.sub_title("4-2. 안전성 (가동범위 준수)")
    pdf.add_table(["지표", "설명", "단위"], SAFETY_METRICS, [w*0.25, w*0.5, w*0.25])

    pdf.sub_title("4-3. 수행 완료도")
    pdf.add_table(["지표", "설명", "값"], COMPLETION_METRICS, [w*0.2, w*0.4, w*0.4])

    # 5. 시계열 데이터
    pdf.add_page()
    pdf.section_title("5. 시계열 데이터 (그래프용)")
    pdf.body_text("평가 중 일정 간격으로 기록되는 스냅샷:")
    pdf.add_table(["항목", "설명"], TIMELINE_METRICS, [w*0.35, w*0.65])
    pdf.note_text("CSV 저장 시에는 경과시간 + 좌/우 유사도만 저장 (경량화)")

    # 6. 데이터 계층 구조
    pdf.section_title("6. 데이터 계층 구조")
    pdf.set_font("malgun", "", 9)
    pdf.set_text_color(0, 0, 0)
    for line in HIERARCHY_TEXT.split("\n"):
        pdf.cell(0, 5, line, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    # 7. CSV 저장 형식
    pdf.section_title("7. CSV 저장 형식")
    pdf.body_text("시나리오 완료 시 자동으로 2개의 CSV 파일 생성:")

    pdf.sub_title("파일명 규칙")
    pdf.body_text("{날짜}_{시간}_{사용자}_{시나리오}_{세션ID}_{종류}.csv")
    pdf.body_text("예: 20260403_153000_홍길동_상부승모근_a1b2c3d4_summary.csv")

    pdf.sub_title("요약 CSV (summary)")
    pdf.body_text("세션 정보 + Phase/Step별 1행. 34개 컬럼.")

    pdf.sub_title("시계열 CSV (timeline)")
    pdf.body_text("시간축 유사도 변화 (웹뷰 그래프용).\n6개 컬럼: SessionId, Phase, Step, Time, LeftSimilarity, RightSimilarity")

    pdf.sub_title("저장 위치")
    pdf.body_text("  - 에디터: Assets/Results/\n  - 빌드 (Quest): Application.persistentDataPath/Results/")

    path = os.path.join(OUTPUT_DIR, "GuideChuna_평가지표_정의서.pdf")
    pdf.output(path)
    print(f"PDF saved: {path}")
    return path


if __name__ == "__main__":
    docx_path = generate_docx()
    pdf_path = generate_pdf()
    print(f"\nDone!\n  DOCX: {docx_path}\n  PDF:  {pdf_path}")
